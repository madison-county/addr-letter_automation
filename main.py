from docx import Document
from docx.shared import Inches
import openpyxl as pyxl

excel_sheet = 'automation_test.xlsx'

def main():

    document = Document()

    wb = pyxl.load_workbook(excel_sheet)
    sheet = wb.active

    excel_parse(sheet)

    section = document.sections[0]
    header = section.header

    paragraph = header.paragraphs[0]
    paragraph.text = 'Madison County GIS Office'
    paragraph.style = document.styles['Header']

    #document.add_heading('Madison County GIS Office', 0)

    p = document.add_paragraph('Lorem ipsum')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic').italic = True

    #document.add_heading('Heading, level 1', level=0)

    document.add_page_break()

    document.save('demo.docx')



def excel_parse(sheet):
    cells = sheet['A2' : 'E3']
    for c1, c2, c3, c4, c5 in cells:
        print(f'{c1.value} - {c2.value} - {c3.value}\n{c4.value}\n{c5.value}\n')


if __name__ == '__main__':
    main()